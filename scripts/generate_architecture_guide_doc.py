#!/usr/bin/env python3
"""Generate ECS Architecture and Deployment Guide (.md + .docx). Documentation only."""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
MD_PATH = DOCS / "ECS_Architecture_and_Deployment_Guide.md"
DOCX_PATH = DOCS / "ECS_Architecture_and_Deployment_Guide.docx"

TODAY = date.today().isoformat()


def build_markdown() -> str:
    return f"""# ECS Architecture and Deployment Guide

**Repository:** `wkin_ecs_consolidated_demo_v13`  
**Document version:** 1.0  
**Last updated:** {TODAY}  
**Classification:** Internal — Architecture Review / Engineering Handover

---

## Table of Contents

1. [Executive Summary](#1-executive-summary)
2. [ECS Overview](#2-ecs-overview)
3. [Business Purpose](#3-business-purpose)
4. [Technology Stack](#4-technology-stack)
5. [High-Level Architecture Diagram](#5-high-level-architecture-diagram)
6. [Component Architecture](#6-component-architecture)
7. [Module Architecture](#7-module-architecture)
8. [Folder Structure](#8-folder-structure)
9. [Database Architecture](#9-database-architecture)
10. [API Architecture](#10-api-architecture)
11. [Security Model](#11-security-model)
12. [Role-Based Access Model](#12-role-based-access-model)
13. [Local Development Setup](#13-local-development-setup)
14. [Deployment Architecture](#14-deployment-architecture)
15. [Dockerization Strategy](#15-dockerization-strategy)
16. [docker-compose Architecture](#16-docker-compose-architecture)
17. [Environment Variables](#17-environment-variables)
18. [Team Handover Guide](#18-team-handover-guide)
19. [Troubleshooting Guide](#19-troubleshooting-guide)
20. [Future Roadmap](#20-future-roadmap)

---

## 1. Executive Summary

The **Enterprise Compliance System (ECS)** consolidated demo (v13) is a single-process **FastAPI + Jinja2** banking governance platform. It delivers role-aware dashboards, framework catalogs, evidence workflows, executive analytics, enterprise GRC surfaces, and AI SDLC governance using **deterministic in-memory mock data**.

This guide documents:

- **As-built architecture** in the repository today (demo-ready, no production database).
- **Target enterprise stack** (React, Node.js, PostgreSQL) for teams planning production hardening.
- **Deployment and handover** guidance for CIO stakeholders, Architecture Review Board (ARB), and engineering teams.

| Dimension | As-built (demo) | Target (enterprise) |
|-----------|-----------------|---------------------|
| UI | Server-rendered HTML + Bootstrap 5 | React + TypeScript + Vite SPA |
| API | FastAPI (Python) | Node.js + Express (or retain FastAPI behind BFF) |
| Data | In-memory Python structures | PostgreSQL |
| Auth | Role picker login (no IdP) | SSO / OIDC + session or JWT |
| Deploy | `uvicorn` local / VM | Docker + orchestration |

---

## 2. ECS Overview

ECS unifies **Governance, Risk, Compliance, Audit, and Evidence Management** into one navigable control surface. The consolidated demo replaces fragmented spreadsheets, ticket queues, and ad-hoc evidence folders with:

- Executive dashboards and trends
- Per-framework control and evidence catalogs (15 frameworks, 307 controls)
- Evidence submission, review, rejection, and closure workflows
- Audit preparation, lifecycle, completeness, and gap analysis
- Enterprise GRC (risk register, CMDB, heatmaps, correlation)
- AI SDLC governance (control tower, workflows, controlled documents)
- Universal KPI/chart drill-down modals

**Demo anchor date:** `2026-05-28` (`DEMO_ANCHOR_DATE` in mock engines).  
**Process model:** One Uvicorn worker; state re-seeded on startup.

---

## 3. Business Purpose

### 3.1 Objectives

1. **Centralize governance data** — one filter engine and catalog across modules.
2. **Continuous audit readiness** — calendars, readiness scores, audit prep cockpit.
3. **Evidence reuse** — cross-framework control theming and overlap intelligence.
4. **Framework onboarding** — loader, mapping, reuse scoring, activation workflows.
5. **Role-tailored experiences** — distinct KPIs, queues, and permissions per persona.
6. **AI-assisted governance** — copilot Q&A, AI governance posture, AI SDLC controls (heuristic demo; no live LLM by default).
7. **Integration narrative** — mocked ServiceNow, SharePoint, Jira, Prisma, SIEM feeds into a single repository story.

### 3.2 Primary audiences

| Audience | Typical use |
|----------|-------------|
| CIO / CISO | Executive dashboards, trends, heatmaps, approvals |
| App Owner | Framework compliance, evidence upload, work queues |
| Auditor | Review queues, framework administration, audit prep |
| Vertical Head | Vertical-scoped compliance and risk views |
| Compliance / Functional heads | Compliance dashboards and analytics |
| ARB / Engineering | Architecture, deployment, module ownership |

---

## 4. Technology Stack

### 4.1 As-built stack (repository truth — `wkin_ecs_consolidated_demo_v13`)

#### Frontend (current)

| Technology | Role |
|------------|------|
| **Jinja2** | Server-side HTML templates |
| **Bootstrap 5.3** | Layout, modals, components (CDN: jsDelivr) |
| **Vanilla JavaScript** | `fetch()` for drill-down APIs, filter clients |
| **Custom CSS** | `enterprise_theme.html`, `executive_charts_system.html` |
| **CSS bar charts** | Executive compact charts (no Chart.js / React) |

Static assets: `modules/shared/static/` mounted at `/static/ecs`.

#### Backend (current)

| Technology | Role |
|------------|------|
| **Python 3.12+** | Runtime |
| **FastAPI** | HTTP routing, JSON APIs, form posts |
| **Uvicorn** | ASGI server (`uvicorn app.main:app`) |
| **Jinja2** | Template rendering via `ChoiceLoader` |
| **python-multipart** | Form uploads |

Dependencies: `requirements.txt` — `fastapi`, `uvicorn`, `jinja2`, `python-multipart`.

#### Data (current)

| Technology | Role |
|------------|------|
| **In-memory structures** | `ecs_state`, catalog constants, ring buffers |
| **No ORM / no SQL** | State resets on process restart |

### 4.2 Target enterprise stack (reference architecture)

The following stack is the **recommended target** for production ECS (not fully implemented in this demo repository). ARB and product teams should treat this as the migration north star.

#### Frontend (target)

| Technology | Purpose |
|------------|---------|
| **React** | Component-based SPA for dashboards and workflows |
| **TypeScript** | Type-safe UI contracts aligned with API schemas |
| **Vite** | Fast dev server, optimized production bundles |

#### Backend (target)

| Technology | Purpose |
|------------|---------|
| **Node.js** | Runtime for API tier or BFF |
| **Express** | REST routing, middleware, integration adapters |

*Alternative:* retain **FastAPI** as the core API and use Node only as a BFF/edge layer.

#### Database (target)

| Technology | Purpose |
|------------|---------|
| **PostgreSQL** | Workflow state, audit trail, evidence metadata, tenancy |

---

## 5. High-Level Architecture Diagram

### 5.1 As-built (demo)

```mermaid
flowchart TB
    subgraph Client
        B[Browser]
    end
    subgraph App["FastAPI + Uvicorn :8000"]
        M[app/main.py]
        R[modules/shared/routes/routes_mvp.py]
        E[Domain Engines]
    end
    subgraph State
        S[(In-Memory ecs_state)]
        C[Framework Catalog]
    end
    B -->|HTML + JSON| M
    B --> R
    M --> E
    R --> E
    E --> S
    E --> C
```

### 5.2 Target (enterprise)

```mermaid
flowchart TB
    subgraph Client
        SPA[React + Vite SPA]
    end
    subgraph API
        NODE[Node.js + Express]
        PY[Optional FastAPI services]
    end
    subgraph Data
        PG[(PostgreSQL)]
        S3[Object Storage]
    end
    SPA --> NODE
    NODE --> PG
    NODE --> S3
    NODE --> PY
```

### 5.3 ASCII — request flow (as-built)

```
[Browser] --HTTP--> [Uvicorn :8000]
                         |
            +------------+------------+
            |            |            |
      app/main.py   routes_mvp   evidence/ai/grc routes
            |            |            |
            +------------+------------+
                         |
              Domain engines (modules/*)
                         |
              In-memory state + catalog
```

---

## 6. Component Architecture

| Layer | Components | Responsibility |
|-------|------------|----------------|
| **Presentation** | Jinja templates, partials, Bootstrap, inline JS clients | Render pages, modals, charts |
| **Routing** | `app/main.py`, `routes_mvp.py`, `evidence_routes.py`, AI/GRC route registrars | HTTP endpoints |
| **Application** | Domain engines per module | Business logic, mock generation |
| **Shared kernel** | `ecs_state`, `enterprise_context`, `role_permissions`, universal drill | Cross-cutting state and UX |
| **Integration (mock)** | `integrations_module`, scheduler, upload | Demo integration story |
| **Observability** | `ecs_logging.py` | Structured console logging |

**Startup sequence (lifespan):**

1. Configure logging  
2. `refresh_repository_from_frameworks()`  
3. `seed_demo_workflow_state()`  
4. `self_heal_governance()`  
5. Serve on port **8000**

---

## 7. Module Architecture

### 7.1 Executive Overview

| Item | Detail |
|------|--------|
| **Path** | `modules/executive_overview/` |
| **Engines** | `executive_analytics_engine`, `reporting_module`, `reports_analytics_engine`, `demo_metrics`, `ecs_reports_engine` |
| **Routes** | `/dashboard`, `/dashboard/cio`, `/mvp/enterprise`, `/mvp/pan-india`, `/mvp/reports`, `/mvp/trends`, `/mvp/demo-overview` |
| **Purpose** | Role dashboards, executive KPIs, trends granularity, report exports |

### 7.2 Frameworks

| Item | Detail |
|------|--------|
| **Path** | `modules/frameworks/` |
| **Engines** | `framework_catalog`, `framework_dashboards`, `framework_onboarding_engine`, KPI/row drill engines |
| **Routes** | `/framework/{{name}}`, `/mvp/framework-loader`, `/mvp/framework-admin` |
| **Purpose** | 15 frameworks, controls, evidence catalog, per-framework dashboards |

### 7.3 Operations

| Item | Detail |
|------|--------|
| **Path** | `modules/operations/` |
| **Engines** | `scheduler_module`, `evidence_repository`, `integrations_module`, `onboarding_engine`, AI Ops |
| **Routes** | `/mvp/scheduler`, `/mvp/upload`, `/mvp/integrations`, `/mvp/onboarding`, AI Ops assistant |
| **Purpose** | Scheduling, bulk upload, integrations, application onboarding |

### 7.4 Governance

| Item | Detail |
|------|--------|
| **Path** | `modules/governance/` |
| **Engines** | `audit_schedule_engine`, `workflow_module`, `governance_intelligence`, `trends_drill_engine`, gap export |
| **Routes** | `/mvp/audit-prep`, `/mvp/evidence-health`, `/mvp/lifecycle`, `/mvp/comparison`, `/evidence/review` |
| **Purpose** | Audit readiness, evidence health, lifecycle, completeness, approval analytics |

### 7.5 Enterprise GRC

| Item | Detail |
|------|--------|
| **Path** | `modules/enterprise_grc/` |
| **Engines** | `enterprise_grc`, `correlation_engine`, `ecs_governance_qa_engine` |
| **Routes** | `/mvp/risk-register`, `/mvp/cmdb`, `/mvp/heatmaps`, `/mvp/governance-analytics`, etc. |
| **Purpose** | Risk register, assets, heatmaps, regulatory mapping, certification reports |

### 7.6 AI SDLC Governance

| Item | Detail |
|------|--------|
| **Path** | `modules/ai_sdlc/` |
| **Engines** | Control tower, workflow store, evidence governance, controlled documents |
| **Routes** | Registered via `app/routes_ai_sdlc_governance.py` → `/mvp/ai-sdlc`, governance surfaces |
| **Purpose** | AI SDLC posture, prompts, stage dashboards, SDLC workflows |

### 7.7 Shared platform

| Item | Detail |
|------|--------|
| **Path** | `modules/shared/` |
| **Services** | `persona_display`, `role_permissions`, `ecs_nav_framework`, `module_capabilities` |
| **Drilldowns** | Universal drill engine, `drilldown_engine.js` |
| **Templates** | Sidebar, UX macros, chart system, copilot dock |

---

## 8. Folder Structure

Actual repository layout (top levels):

```
wkin_ecs_consolidated_demo_v13/
├── app/                          # FastAPI entry + compatibility shims → modules/*
│   ├── main.py                   # Application factory, lifespan, core routes
│   ├── ecs_state.py              # Shim → modules/shared/services/ecs_state.py
│   ├── routes_mvp.py             # Shim → modules/shared/routes/routes_mvp.py
│   ├── evidence_routes.py
│   ├── routes_ai_sdlc_governance.py
│   ├── routes_grc_demo.py
│   └── *.py                      # Re-export shims for engines
├── modules/
│   ├── shared/                   # routes, services, templates, static, drilldowns
│   ├── executive_overview/       # engines + templates
│   ├── frameworks/
│   ├── governance/
│   ├── operations/
│   ├── enterprise_grc/           # engines + reports/
│   └── ai_sdlc/
├── tests/                        # pytest (23+ test modules)
├── scripts/                      # validation, certification, doc generation
├── docs/                         # architecture and migration docs
├── requirements.txt
├── start_ecs.sh
├── ECS_ARCHITECTURE_BASELINE.md
└── README.md
```

**Convention:** New logic belongs under `modules/<domain>/engines/`; `app/<module>.py` shims preserve legacy imports.

---

## 9. Database Architecture

### 9.1 As-built

There is **no database**. Persistence uses:

| Store | Location | Contents |
|-------|----------|----------|
| Workflow state | `ecs_state` | Submissions, rejections, clarifications |
| Catalog | `framework_catalog.py` | Frameworks, controls, evidence definitions |
| Audit trail | `audit_trail.py` | Ring buffer (~200 events) |
| Notifications | `audit_trail.py` | Ring buffer (~30 events) |

**Implications:** Restart clears mutable workflow unless re-seeded; no multi-tenant isolation; no concurrent write safety.

### 9.2 Target (PostgreSQL)

Recommended schemas (illustrative):

| Schema area | Tables (examples) |
|-------------|-------------------|
| `core` | `tenants`, `users`, `roles` |
| `governance` | `frameworks`, `controls`, `evidence_items`, `observations` |
| `workflow` | `transitions`, `approvals`, `audit_events` |
| `analytics` | `kpi_snapshots`, `trend_series` |

Migration path: mirror in-memory dicts first (SQLite/Postgres), then externalize file blobs to object storage.

---

## 10. API Architecture

### 10.1 Route registration

| Registrar | File | Scope |
|-----------|------|-------|
| Core | `app/main.py` | Login, dashboards, framework pages, evidence review POSTs |
| MVP | `modules/shared/routes/routes_mvp.py` | `/mvp/*`, most `/api/*` |
| Evidence | `app/evidence_routes.py` | Evidence upload/submit APIs |
| AI SDLC | `app/routes_ai_sdlc_governance.py` | AI SDLC module |
| GRC demo | `app/routes_grc_demo.py` | GRC demo endpoints |

### 10.2 API categories

| Category | Examples | Response |
|----------|----------|----------|
| **Universal drill** | `GET /api/ecs/universal-drill` | JSON modal payload + rows |
| **Framework drill** | `GET /api/framework/kpi-drill` | JSON |
| **Audit prep** | `GET /api/audit-prep/kpi-drill` | JSON |
| **Demo engine** | `GET /api/demo/*` | JSON metrics |
| **Filters** | `GET/POST /api/ecs/filters/*` | JSON filter state |
| **Workflow** | `GET /api/evidence-workflow/summary` | JSON |

### 10.3 Page state contract

UI state is carried in **URL query parameters**: `role`, `user`, `framework`, filter keys. No client-side router (as-built).

---

## 11. Security Model

### 11.1 As-built (demo)

| Control | Status |
|---------|--------|
| Authentication | Role selected on login form only |
| Session / JWT | Not implemented |
| Authorization | Query-param role + `role_permissions.py` |
| CSRF | Not systematically enforced |
| Secrets | None in repo |
| Data encryption | N/A (in-memory) |
| Audit logging | In-memory ring buffer |

**Risk posture:** Suitable for **isolated demo environments only**. Do not expose to untrusted networks without hardening.

### 11.2 Target (production)

| Control | Recommendation |
|---------|----------------|
| Identity | Enterprise IdP (OIDC/SAML) |
| Session | HttpOnly cookies or short-lived JWT |
| Authorization | Server-side RBAC/ABAC mapped to PostgreSQL roles |
| Transport | TLS 1.2+ everywhere |
| Secrets | Vault / K8s secrets, `.env` not committed |
| Evidence files | Virus scan + encrypted object storage |

---

## 12. Role-Based Access Model

Login: `POST /login` with `role` field → redirect with `?role=…&user=…`.

### 12.1 Primary personas

| Role key | Display persona (example) | Default landing |
|----------|---------------------------|-----------------|
| **cio** | R. Khanna — Chief Information Officer | `/dashboard/cio?role=cio&user=CIO` |
| **owner** | R. Sharma — Application Owner | `/dashboard?role=owner&user=AppOwner` |
| **auditor** | A. Banerjee — Lead Auditor | `/dashboard?role=auditor&user=Auditor` |
| **vertical_head** | K. Reddy — Vertical Head | `/dashboard/vertical-head?role=vertical_head&user=VerticalHead` |

### 12.2 Additional roles

| Role key | Default user token | Landing |
|----------|-------------------|---------|
| compliance_head / compliance_officer | ComplianceOfficer | `/dashboard/compliance-head` |
| functional_head | FunctionalHead | `/dashboard/functional-head` |
| security_officer | SecurityOfficer | Compliance dashboard |
| operations_owner | OpsOwner | `/mvp/onboarding` |
| ai_governance_owner | AIGovOwner | AI governance |
| ai_sdlc_owner | SDLCOwner | `/mvp/ai-sdlc` |
| framework_owner | FrameworkOwner | `/mvp/framework-admin` |

### 12.3 Data scoping

`role_filter_scope.py` filters mock datasets (applications, frameworks) per role. **CIO and Auditor** typically see enterprise-wide data; **App Owner** sees owned applications; **Vertical Head** sees vertical scope.

Permissions surface in templates as `perm_*` flags from `permission_ctx()`.

---

## 13. Local Development Setup

### 13.1 Prerequisites

| Requirement | As-built demo | Target stack |
|-------------|---------------|--------------|
| **Runtime** | Python 3.12+ | Node 20 LTS + Python (optional) |
| **Node version** | N/A (no Node in repo) | **20.x LTS** (recommended) |
| **PostgreSQL** | N/A | **16.x** (recommended) |
| **Package manager** | pip | pip + npm/pnpm |
| **Network** | CDN access for Bootstrap | Same + npm registry |

### 13.2 Installation steps (as-built)

```bash
cd wkin_ecs_consolidated_demo_v13
python3 -m venv .venv
source .venv/bin/activate
pip install -r requirements.txt
pip install pytest   # optional
```

### 13.3 Startup commands

| Component | Command |
|-----------|---------|
| **Backend** | `uvicorn app.main:app --reload --host 127.0.0.1 --port 8000` |
| **Convenience** | `./start_ecs.sh` |
| **Frontend (as-built)** | *Not applicable — SSR* |
| **Frontend (target)** | `npm run dev` (Vite, when introduced) |
| **Database (as-built)** | *Not applicable* |
| **Database (target)** | `docker compose up -d postgres` |

**Application URL:** http://127.0.0.1:8000/

### 13.4 Verification

1. Open `/` and log in as each primary role.  
2. Confirm sidebar persona and landing route.  
3. Open `/mvp/trends` and drill a chart bar (modal + row count).

---

## 14. Deployment Architecture

### 14.1 Demo deployment

```
[Developer laptop / VM]
        |
   Uvicorn :8000
        |
   In-memory state
```

### 14.2 Production deployment (target)

```
[Users] → [TLS LB] → [Web tier: static SPA or CDN]
                    → [API tier: Node/Express or FastAPI]
                    → [PostgreSQL primary + replica]
                    → [Object storage for evidence]
                    → [Optional: Redis queue, SIEM]
```

| Environment | Purpose | Notes |
|-------------|---------|-------|
| DEV | Engineer laptops | `--reload`, in-memory or local Postgres |
| UAT | Business validation | Seeded Postgres, masked data |
| PROD | Live | HA Postgres, secrets manager, WAF |

---

## 15. Dockerization Strategy

**Status:** Not present in repository today. Recommended approach:

### 15.1 Images

| Image | Base | Contents |
|-------|------|----------|
| `ecs-api` | `python:3.12-slim` | FastAPI app (current) or `node:20-alpine` (target API) |
| `ecs-web` | `nginx:alpine` | Built Vite `dist/` (target) |
| `ecs-migrate` | API base | DB migration runner (Alembic/Flyway) |

### 15.2 Practices

- Multi-stage builds; non-root user  
- Health check: `GET /` or `/health`  
- Read-only root filesystem where possible  
- Pin dependency versions in image build  

---

## 16. docker-compose Architecture

Illustrative **target** compose stack (for local/full-stack dev):

```yaml
# Illustrative — not committed to repo
services:
  postgres:
    image: postgres:16
    ports: ["5432:5432"]
    environment:
      POSTGRES_DB: ecs
      POSTGRES_USER: ecs
      POSTGRES_PASSWORD: ecs_dev
    volumes: [pgdata:/var/lib/postgresql/data]

  api:
    build: ./docker/api
    ports: ["8000:8000"]
    env_file: .env
    depends_on: [postgres]

  web:
    build: ./docker/web
    ports: ["5173:80"]
    depends_on: [api]

volumes:
  pgdata:
```

**As-built demo compose (minimal):**

```yaml
services:
  ecs-demo:
    build: .
    command: uvicorn app.main:app --host 0.0.0.0 --port 8000
    ports: ["8000:8000"]
```

---

## 17. Environment Variables

### 17.1 As-built

No environment variables are required. Configuration is code-defined (`DEMO_MODE`, `DEMO_ANCHOR_DATE`).

### 17.2 `.env.example` (target / recommended)

```bash
# ECS Environment — example only (not required for current demo)

# Application
APP_ENV=development
APP_HOST=0.0.0.0
APP_PORT=8000
LOG_LEVEL=INFO
DEMO_MODE=true
DEMO_ANCHOR_DATE=2026-05-28

# Database (target)
DATABASE_URL=postgresql://ecs:ecs_dev@localhost:5432/ecs
DATABASE_POOL_SIZE=10

# Node API (target)
NODE_ENV=development
API_PORT=4000
CORS_ORIGIN=http://localhost:5173

# Auth (target)
OIDC_ISSUER=https://login.example.com/
OIDC_CLIENT_ID=ecs-app
OIDC_CLIENT_SECRET=change-me

# Object storage (target)
S3_BUCKET=ecs-evidence-dev
S3_REGION=ap-south-1

# Integrations (target)
SERVICENOW_INSTANCE=
SHAREPOINT_TENANT_ID=
```

Copy to `.env` locally; **never commit** secrets.

---

## 18. Team Handover Guide

### 18.1 First-week checklist

1. Read `ECS_ARCHITECTURE_BASELINE.md` and `docs/ECS_DEPENDENCY_REPORT.md`.  
2. Run `uvicorn app.main:app --reload` and walk through four primary roles.  
3. Run `.venv/bin/python -m pytest tests/test_platform_ui.py -q`.  
4. Trace one drill-down: UI click → `/api/ecs/universal-drill` → drill engine.  

### 18.2 Ownership

| Change type | Primary touchpoints |
|-------------|---------------------|
| New MVP page | `routes_mvp.py`, template, engine |
| Persona / header | `persona_display.py`, `enterprise_context.py` |
| Permissions | `role_permissions.py`, `role_filter_scope.py` |
| Framework data | `framework_catalog.py` |
| Charts / trends | `executive_analytics_engine.py`, `executive_charts_system.html` |

### 18.3 Shim pattern

Imports may use `app.framework_catalog` or `modules.frameworks.engines.framework_catalog` — both resolve via shims. Prefer **`modules.*`** for new code.

### 18.4 Certification artifacts

Reports under `modules/enterprise_grc/reports/` (route matrix, platform certification JSON/CSV).

---

## 19. Troubleshooting Guide

| Symptom | Likely cause | Action |
|---------|--------------|--------|
| Port 8000 in use | Prior Uvicorn | `pkill -f uvicorn` or change port |
| Unstyled pages | CDN blocked | Allow jsDelivr or vendor Bootstrap locally |
| Empty drill modal | Missing `role` in API call | Ensure `drilldown_engine.js` passes `role` |
| State “lost” after restart | In-memory design | Expected; lifespan re-seeds demo |
| Template error after edit | Jinja strictness | Run `python scripts/validate_templates.py` |
| Import errors | Wrong path | Use module path or existing shim |
| pytest not found | No venv | `pip install pytest` in `.venv` |

---

## 20. Future Roadmap

| Phase | Initiative | Outcome |
|-------|------------|---------|
| **1** | PostgreSQL persistence | Durable workflow and audit trail |
| **2** | OIDC authentication | Enterprise login |
| **3** | React + Vite frontend | SPA with typed API client |
| **4** | Node/Express BFF (optional) | Integration orchestration |
| **5** | Docker + CI/CD | Repeatable deploys |
| **6** | Real integrations | ServiceNow, SharePoint, SIEM connectors |
| **7** | LLM governance | Policy-bound copilot with audit |
| **8** | Multi-tenant | Bank / entity isolation |

---

## Document control

| Version | Date | Author | Notes |
|---------|------|--------|-------|
| 1.0 | {TODAY} | ECS Platform Team | Initial consolidated guide from repository analysis |

---

*End of document*
"""


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char2)
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char3)


def set_heading_style(paragraph, level: int):
    paragraph.style = f"Heading {level}"


def _flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row_cells in enumerate(rows):
        for ci, val in enumerate(row_cells):
            table.rows[ri].cells[ci].text = val


def md_to_docx(md: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    table_buf: list[list[str]] = []

    def flush_tables():
        nonlocal table_buf
        if table_buf:
            _flush_table(doc, table_buf)
            table_buf = []

    # Title page
    title = doc.add_paragraph()
    title.alignment = 1  # center
    run = title.add_run("ECS Architecture and Deployment Guide")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = 1
    sub_run = sub.add_run(f"wkin_ecs_consolidated_demo_v13\nVersion 1.0 — {TODAY}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()
    meta = doc.add_paragraph(
        "Audience: CIO, Architecture Review Board, Development Team\n"
        "Classification: Internal"
    )
    meta.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_page_break()

    # TOC
    h = doc.add_heading("Table of Contents", level=1)
    add_toc(doc.add_paragraph())
    doc.add_paragraph("(Update fields in Word: right-click Table of Contents → Update Field)")
    doc.add_page_break()

    in_code = False
    code_lines: list[str] = []
    skip_until_toc = True

    for line in md.splitlines():
        if skip_until_toc:
            if line.startswith("## 1. Executive Summary"):
                skip_until_toc = False
            else:
                continue

        if line.startswith("```"):
            flush_tables()
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("## "):
            flush_tables()
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            flush_tables()
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            flush_tables()
            doc.add_heading(line[5:].strip(), level=3)
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            continue

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):
                continue
            table_buf.append(cells)
            continue

        flush_tables()

        if line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
            continue

        if not line.strip():
            continue

        # bold inline **text**
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2])
                r.bold = True
            else:
                p.add_run(part)

    flush_tables()
    return doc


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    md = build_markdown()
    MD_PATH.write_text(md, encoding="utf-8")
    doc = md_to_docx(md)
    doc.save(DOCX_PATH)
    print(f"Wrote {MD_PATH}")
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
